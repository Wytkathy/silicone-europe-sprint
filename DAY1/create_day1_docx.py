from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "市场分析入门课.md"
OUTPUT = ROOT / "DAY1_制造工厂市场分析入门课.docx"

BLUE = "1D5D4C"
DARK = "102F2A"
INK = "172B27"
MUTED = "66756F"
LIGHT = "E5F1E9"
PALE = "F5F1E8"
GRAY = "F2F4F3"
WHITE = "FFFFFF"
CLAY = "C96E55"


def set_font(run, size=10.5, bold=None, color=INK, italic=None, mono=False):
    latin = "Menlo" if mono else "Aptos"
    east = "PingFang SC"
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), latin)
    rpr.rFonts.set(qn("w:hAnsi"), latin)
    rpr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), color)


def cell_width(cell, width):
    props = cell._tc.get_or_add_tcPr()
    tcw = props.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        props.append(tcw)
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")


def cell_margins(cell):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for key, value in (("top", 100), ("start", 120), ("bottom", 100), ("end", 120)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    props = table._tbl.tblPr
    tblw = props.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        props.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell_width(cell, widths[idx])
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def rich(paragraph, value, size=10.5, color=INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", value)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size - 0.5, color=DARK, mono=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size, color=color, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DAY 1  ·  制造工厂市场分析入门课")
    set_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Silora Europe Sprint  |  ")
    set_font(run, size=8.5, color=MUTED)
    field_run = footer.add_run()
    for kind, value in [("begin", None), (None, "PAGE"), ("separate", None), ("end", None)]:
        node = OxmlElement("w:instrText" if value else "w:fldChar")
        if value:
            node.set(qn("xml:space"), "preserve")
            node.text = value
        else:
            node.set(qn("w:fldCharType"), kind)
        field_run._r.append(node)


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("DAY 1 · MARKET ANALYSIS")
    set_font(run, size=11, bold=True, color=CLAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("制造工厂的市场分析入门课")
    set_font(run, size=27, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("从市场假设、客户开发到 AI 工作流与 OBM 小实验")
    set_font(run, size=13, color=BLUE)

    table = doc.add_table(rows=4, cols=2)
    table_geometry(table, [2100, 7260])
    values = [
        ("适合对象", "即将进入 ODM/OEM 制造工厂、没有市场分析经验的学习者"),
        ("核心任务", "学会用证据降低产品、客户和市场决策的不确定性"),
        ("案例方向", "德国设计型厨房市场与可重复使用硅胶食品收纳"),
        ("学习方式", "阅读 → 提出假设 → 收集证据 → 低成本验证 → 复盘"),
    ]
    for idx, (label, value) in enumerate(values):
        shade(table.cell(idx, 0), LIGHT)
        shade(table.cell(idx, 1), PALE)
        p1 = table.cell(idx, 0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.space_after = Pt(0)
        r = p1.add_run(label)
        set_font(r, size=9.5, bold=True, color=DARK)
        p2 = table.cell(idx, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(value)
        set_font(r, size=9.5)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def code_box(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    shade(table.cell(0, 0), "F4F7F5")
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, size=8.8, color=DARK, mono=True)


def note_box(doc, value):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    shade(table.cell(0, 0), LIGHT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    rich(p, value, size=10.5, color=DARK)


def new_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "7")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def widths_for(cols):
    if cols == 2:
        return [2500, 6860]
    if cols == 3:
        return [1900, 3560, 3900]
    if cols == 4:
        return [1600, 2500, 2560, 2700]
    return [9360 // cols] * cols


def markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r"[-: ]+", cell or "-") for cell in cells):
            continue
        parsed.append(cells)
    cols = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=cols)
    table_geometry(table, widths_for(cols))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_idx, row in enumerate(parsed):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                shade(cell, BLUE)
            elif row_idx % 2 == 0:
                shade(cell, GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rich(p, value, size=8.8 if cols >= 3 else 9.2, color=WHITE if row_idx == 0 else INK)
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True


def convert(doc, text):
    lines = text.splitlines()
    idx = 0
    skipped_title = False
    active_numbering = None
    while idx < len(lines):
        value = lines[idx].strip()
        if not value or value == "---":
            active_numbering = None
            idx += 1
            continue
        if value.startswith("# ") and not skipped_title:
            skipped_title = True
            active_numbering = None
            idx += 1
            continue
        if value.startswith("> "):
            active_numbering = None
            note_box(doc, value[2:])
            idx += 1
            continue
        if value.startswith("```"):
            active_numbering = None
            block = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                block.append(lines[idx])
                idx += 1
            code_box(doc, block)
            idx += 1
            continue
        if value.startswith("|"):
            active_numbering = None
            rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append(lines[idx])
                idx += 1
            markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", value)
        if heading:
            active_numbering = None
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            rich(p, heading.group(2), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK)
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", value)
        if bullet:
            active_numbering = None
            p = doc.add_paragraph(style="List Bullet")
            rich(p, bullet.group(1))
            idx += 1
            continue
        number = re.match(r"^\d+\.\s+(.+)$", value)
        if number:
            if active_numbering is None:
                active_numbering = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_numbering)
            rich(p, number.group(1))
            idx += 1
            continue
        active_numbering = None
        parts = [value]
        idx += 1
        while idx < len(lines):
            next_value = lines[idx].strip()
            if (
                not next_value
                or next_value.startswith(("#", "-", "|", "```", ">"))
                or re.match(r"^\d+\.\s+", next_value)
            ):
                break
            parts.append(next_value)
            idx += 1
        p = doc.add_paragraph()
        rich(p, " ".join(parts))


document = Document()
setup(document)
cover(document)
convert(document, SOURCE.read_text(encoding="utf-8"))
document.core_properties.title = "DAY1 制造工厂市场分析入门课"
document.core_properties.subject = "市场分析、客户开发、AI 工作流与 OBM"
document.core_properties.author = "Codex"
document.save(OUTPUT)
print(OUTPUT)
